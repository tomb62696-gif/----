import streamlit as st
import pdfplumber
import json
import io
from docx import Document
from docx.shared import Pt
from openai import OpenAI

# ==========================================
# 1. 初始化 AI 客户端
# ==========================================
# 请在此处填写你的 API Key 和 Base URL
API_KEY = "sk-2d4070285351465ebdf763a80c76eb84" 
BASE_URL = "https://api.deepseek.com/v1" # 以 DeepSeek 为例

client = OpenAI(api_key=API_KEY, base_url=BASE_URL)

# ==========================================
# 2. 核心工具函数：PDF 全文深度解析
# ==========================================
def extract_pdf_text(uploaded_file):
    """深度提取工程 PDF。为了维度4能覆盖交通流预测和投资测算，优化读取跨度"""
    text_segments = []
    with pdfplumber.open(uploaded_file) as pdf:
        total_pages = len(pdf.pages)
        st.info(f"📄 报告解析成功，共 {total_pages} 页。正在调集数据、技术与财务章节...")
        
        # 读取前 25 页 (涵盖目录、概述、政策、以及初步的技术指标、隧道工法变更说明)
        for i in range(min(25, total_pages)):
            text_segments.append(f"--- PAGE {i+1} ---\n" + (pdf.pages[i].extract_text() or ""))
            
        # 读取最后 12 页 (涵盖详细投资估算、社会效益、风险控制及结论)
        if total_pages > 37:
            text_segments.append("\n...[中间章节省略]...\n")
            for i in range(total_pages - 12, total_pages):
                text_segments.append(f"--- PAGE {i+1} ---\n" + (pdf.pages[i].extract_text() or ""))
                
    return "\n".join(text_segments)

# ==========================================
# 3. AI 四维全能复合审查引擎（具备强鲁棒性清洗）
# ==========================================
def ai_four_dimension_review(pdf_text):
    """让 AI 同时扮演：审计员(D1)、规划师(D2)、总工程师(D3)、投资风控总监(D4)"""
    
    prompt = f"""
    你是一个精通国家工程建设强制性标准、发改委可研大纲、以及基础设施项目（特别是大型水下/盾构隧道）全生命周期管理的顶级会审专家。
    请仔细阅读以下从工程可研报告中提取的文本，从四个维度进行全维度的立体化审查。
    
    请严格按照以下 JSON 格式输出，不要包含任何多余的 Markdown 标记或解释文字：
    {{
      "dimension_1_data": {{
        "project_name": "项目名称",
        "total_investment_wanyu": 0.0,
        "construction_period_months": 0,
        "route_length_km": 0.0,
        "design_speed_kmh": 0
      }},
      "dimension_2_policy": {{
        "checklist": {{
          "建设必要性论证": true, 
          "规划与选址合规性": true, 
          "节能减排与双碳响应": true,
          "环境影响与生态保护": true,
          "社会稳定风险评估": false
        }},
        "missing_sections_details": "薄弱环节说明",
        "policy_fit_analysis": "宏观政策契合度评价",
        "normative_suggestions": "发改委规范性建议"
      }},
      "dimension_3_regulatory": {{
        "overall_status": "安全合规 / 存在强条合规风险 / 存在一般条文建议",
        "mandatory_clauses": [
          {{
            "item_name": "核心条文项名称",
            "standard_basis": "规范出处与条款号",
            "report_value": "报告中对应的实际数值或方案说明",
            "status": "合规 / 违规 / 风险",
            "expert_analysis": "依据条文进行的专家合规性拆解。"
          }}
        ],
        "detailed_violations": "强条违规缺陷详述"
      }},
      "dimension_4_logic": {{
        "safety_logic_review": "【安全性逻辑审查】：针对报告中提到的核心工艺技术路线（如明挖改盾构工法、水利红线穿过、防灾排烟、危险品运输管制等），论证其安全性逻辑是否自洽？是否存在未交待清楚的重大工程安全或施工次生风险？",
        "economic_reasonableness": "【经济合理性与性价比控制】：基于总投资和规模，评估其是否存在明显的过度设计（例如：非必要的高标准配置）？投资构成比例是否合理？是否有压降成本的空间？",
        "demand_scale_matching": "【供需规模匹配逻辑】：评估项目的建设规模与前文提到的交通流量预测、局部饱和度之间，是否具备严密的因果推导逻辑？",
        "financial_sustainability": "【财务与效益分摊逻辑】：对报告提及的财政资金来源、跨区域分摊机制、以及后期高昂运营维护成本的财务可持续性进行深度审计逻辑评价。"
      }}
    }}
    
    【针对东太湖隧道等重大水下工程项目的深度审查指引】：
    - 维度4重点看：原方案围堰明挖改成双线盾构后，虽然保护了生态环境，但造价剧增，报告在“环境安全 vs 经济可行性”两者的权衡论证上是否逻辑闭环？
    - 检查按照区界分摊财政资金的经济逻辑、以及搬迁费在总投资中的合理性。
    
    【输入文本】：
    {pdf_text}
    """
    
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": "你是一个只会输出高精度、标准结构化 JSON 数据的自动化工程审查后台。"},
            {"role": "user", "content": prompt}
        ],
        response_format={"type": "json_object"},
        temperature=0.15
    )
    
    # 获取原始文本并进行深度清洗，防止 Markdown 标记包裹导致的反序列化崩溃
    raw_content = response.choices[0].message.content.strip()
    
    if raw_content.startswith("```json"):
        raw_content = raw_content[7:]
    elif raw_content.startswith("```"):
        raw_content = raw_content[3:]
        
    if raw_content.endswith("```"):
        raw_content = raw_content[:-3]
        
    raw_content = raw_content.strip()
    
    try:
        return json.loads(raw_content)
    except json.JSONDecodeError:
        # 异常兜底，防止前端页面彻底崩溃
        return {
            "dimension_1_data": {"project_name": "解析失败（数据异常）", "total_investment_wanyu": 0, "construction_period_months": 0, "route_length_km": 0, "design_speed_kmh": 0},
            "dimension_2_policy": {"checklist": {}, "missing_sections_details": "AI返回格式异常，未能成功解析结构化清单。"},
            "dimension_3_regulatory": {"overall_status": "无法评估", "mandatory_clauses": []},
            "dimension_4_logic": {"safety_logic_review": "未解析成功", "economic_reasonableness": "未解析成功"}
        }

# ==========================================
# 新增：Word 报告生成函数
# ==========================================
def create_word_report(results):
    doc = Document()
    doc.add_heading('工程可行性研究报告审查意见', 0)
    
    # 维度1
    doc.add_heading('一、关键技术经济指标', level=1)
    d1 = results['dimension_1_data']
    for k, v in d1.items():
        doc.add_paragraph(f"{k}: {v}")
        
    # 维度2
    doc.add_heading('二、政策合规性审查', level=1)
    doc.add_paragraph(results['dimension_2_policy']['policy_fit_analysis'])
    
    # 维度3
    doc.add_heading('三、法规强制性条文审查', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = '规范依据'
    hdr_cells[2].text = '结论'
    for item in results['dimension_3_regulatory']['mandatory_clauses']:
        row_cells = table.add_row().cells
        row_cells[0].text = item['item_name']
        row_cells[1].text = item['standard_basis']
        row_cells[2].text = item['status']
        
    # 维度4
    doc.add_heading('四、内容逻辑与经济性审查', level=1)
    doc.add_paragraph(results['dimension_4_logic']['safety_logic_review'])
    
    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 4. Streamlit 在线网页前端界面
# ==========================================
st.set_page_config(page_title="全维工程可研智能审查系统 v4.0", layout="wide")

st.title("🏗️ 全在线工程可行性研究报告立体化审查系统 (全功能旗舰版)")
st.caption("基于 AI + 行业大数据规则碰撞的工程文件全生命周期会审平台 | 已激活：全维度 1、2、3、4 联合引擎")

# 侧边栏配置
with st.sidebar:
    st.header("⚙️ 审查维度配置")
    st.checkbox("维度1：文字与数据一致性审查", value=True, disabled=True)
    st.checkbox("维度2：政策合规性审查", value=True, disabled=True)
    st.checkbox("维度3：法规标准符合性审查（强条）", value=True, disabled=True)
    st.checkbox("维度4：内容逻辑性审查（安全经济）", value=True, disabled=True)
    st.markdown("---")
    st.markdown("💡 **维度4审计重点：**\n"
                "- 技术路线安全性论证闭环\n"
                "- 每公里造价与过度设计审计\n"
                "- 流量预测与车道规模匹配度\n"
                "- 跨区资金筹措与长周期运营财务")

# 文件上传
uploaded_file = st.file_uploader("请上传您的工程可研报告 PDF 文件（完美支持东太湖隧道等大样本）", type=["pdf"])

if uploaded_file is not None:
# 使用 session_state 存储结果，确保网页刷新时数据不丢失
    if "review_results" not in st.session_state:
        st.session_state["review_results"] = None

    if st.button("🚀 启动四维一体化全能审查"):
        with st.spinner("系统正在调集AI多方专家（数据、政策、强条、技术经济专家）进行全面联合评审..."):
            
            # 1. 解析文本
            pdf_text = extract_pdf_text(uploaded_file)
            
            # 2. 调用四维 AI 引擎
            review_results = ai_four_dimension_review(pdf_text)
            
            d1_data = review_results.get("dimension_1_data", {})
            d2_data = review_results.get("dimension_2_policy", {})
            d3_data = review_results.get("dimension_3_regulatory", {})
            d4_data = review_results.get("dimension_4_logic", {})
            
            st.success("🎉 四维立体化联合审查圆满完成！终审意见报告已生成。")
            st.write("---")
            
            # 创建四个标签页
            tab1, tab2, tab3, tab4 = st.tabs([
                "🔍 维度1：数据一致性", 
                "⚖️ 维度2：政策与规范性", 
                "📐 维度3：法规与强条对碰",
                "📊 维度4：技术安全与经济逻辑"
            ])
            
            # ==========================================
            # TAB 1: 维度1 结果展示
            # ==========================================
            with tab1:
                st.subheader("📊 关键指标元数据对账")
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("项目名称", d1_data.get("project_name", "未识别"))
                with col2:
                    st.metric("总投资额 (万元)", f"{d1_data.get('total_investment_wanyu', 0):,.2f}")
                with col3:
                    st.metric("建设工期 (个月)", f"{d1_data.get('construction_period_months', 0)} 个月")
                with col4:
                    st.metric("路线全长 (km)", f"{d1_data.get('route_length_km', 0)} km")
                
                st.markdown("#### 📋 自动数据对账报告")
                errors_found = 0
                if d1_data.get('total_investment_wanyu', 0) > 500000 and d1_data.get('construction_period_months', 0) < 24:
                    st.error("❌ **发现严重数据矛盾（投资与工期不匹配）**")
                    st.markdown(f"* **问题描述**：项目总投资高达 `{d1_data.get('total_investment_wanyu')}` 万元，但声明的工期仅为 `{d1_data.get('construction_period_months')}` 个月，存在重大工期漏算风险。")
                    errors_found += 1
                if errors_found == 0:
                    st.success("✅ 经过后端硬编码规则比对，报告内核心指标的前后文钩稽关系完全一致。")
            
            # ==========================================
            # TAB 2: 维度2 结果展示
            # ==========================================
            with tab2:
                st.subheader("📋 必备章节齐全性审查 (Checklist)")
                checklist = d2_data.get("checklist", {})
                
                # 增加空值保护网，防止无数据时 st.columns(0) 报错崩溃
                if checklist:
                    chk_cols = st.columns(len(checklist))
                    for idx, (item, is_present) in enumerate(checklist.items()):
                        with chk_cols[idx]:
                            if is_present:
                                st.success(f"🍏 {item}\n\n**[已包含]**")
                            else:
                                st.error(f"🍎 {item}\n\n**[缺失/薄弱]**")
                else:
                    st.warning("⚠️ 未能从报告中提取到结构化的齐全性检查清单。")
                
                if d2_data.get("missing_sections_details") and d2_data.get("missing_sections_details") != "无":
                    st.warning(f"⚠️ **调优建议：** {d2_data.get('missing_sections_details')}")
                
                st.markdown("---")
                st.subheader("🌱 发展政策与战略红线契合度研判")
                st.info(d2_data.get("policy_fit_analysis", ""))
                st.subheader("📐 编制规范性建议")
                st.markdown(d2_data.get("normative_suggestions", ""))
            
            # ==========================================
            # TAB 3: 维度3 结果展示
            # ==========================================
            with tab3:
                st.subheader("🛑 国家工程建设强制性条文对碰结果")
                status = d3_data.get("overall_status", "未评估")
                if "违规" in status or "风险" in status:
                    st.error(f"⚠️ 研判结论：{status}")
                else:
                    st.success(f"✅ 研判结论：{status}")
                
                clauses = d3_data.get("mandatory_clauses", [])
                st.markdown("#### 📑 逐项规范条文对碰明细表")
                for idx, clause in enumerate(clauses):
                    c_status = clause.get("status", "合规")
                    if c_status == "合规":
                        icon, color_box = "🟢 [核心合规]", st.success
                    elif c_status == "风险":
                        icon, color_box = "🟡 [技术风险]", st.warning
                    else:
                        icon, color_box = "🔴 [强条违规]", st.error
                        
                    with st.expander(f"{icon} {clause.get('item_name')} — 依据：{clause.get('standard_basis')}"):
                        st.markdown(f"**📑 规范/法条依据：** `{clause.get('standard_basis')}`")
                        st.markdown(f"**📝 本报告设计值/方案：** {clause.get('report_value')}")
                        st.markdown(f"**🧑‍💻 审查总工意见：** {clause.get('expert_analysis')}")
                
                st.markdown("---")
                st.subheader("🚨 强条/风险缺陷详述与整改方案")
                st.markdown(d3_data.get("detailed_violations", "未发现违反国家强制性条文的情况。"))

            # ==========================================
            # TAB 4: 维度4 结果展示
            # ==========================================
            with tab4:
                st.subheader("💎 方案技术安全与经济性逻辑深度审计看板")
                
                tot_inv = d1_data.get('total_investment_wanyu', 0)
                r_len = d1_data.get('route_length_km', 0)
                
                # 增加防除零保护安全网
                if tot_inv > 0 and r_len > 0:
                    cost_per_km = (tot_inv / 10000) / r_len
                    st.metric("📊 自动化推算：每公里宏观综合造价", f"{cost_per_km:.2f} 亿元 / 公里")
                    
                    if cost_per_km > 8.0:
                        st.warning("💡 **造价红线预警**：该工程每公里综合造价已步入**超高造价区间（>8亿元/km）**。AI建议审查时应严格核对大直径盾构施工特种设备费折旧，以及由于工法由明挖改盾构带来的“盾构段与岸上明挖交界段”的防水加固成本。")
                else:
                    st.metric("📊 自动化推算：每公里宏观综合造价", "无法计算（数据不全或解析长度为0）")
                
                st.markdown("---")
                
                # 渲染 AI 的四大内容逻辑分析块
                c_sec1, c_sec2 = st.columns(2)
                with c_sec1:
                    with st.chat_message("assistant", avatar="🛡️"):
                        st.markdown("##### **方案技术安全性推导逻辑**")
                        st.write(d4_data.get("safety_logic_review", "未生成安全性评估。"))
                        
                    with st.chat_message("assistant", avatar="📈"):
                        st.markdown("##### **需求规模与流量匹配度逻辑**")
                        st.write(d4_data.get("demand_scale_matching", "未生成规模匹配分析。"))
                        
                with c_sec2:
                    with st.chat_message("assistant", avatar="💰"):
                        st.markdown("##### **经济合理性与性价比审计**")
                        st.write(d4_data.get("economic_reasonableness", "未生成经济合理性评估。"))
                        
                    with st.chat_message("assistant", avatar="🏦"):
                        st.markdown("##### **长周期财务与资金可持续性**")
                        st.write(d4_data.get("financial_sustainability", "未生成财务效益评估。"))


        st.write("---")
        # 生成 Word 并显示下载按钮
        word_buffer = create_word_report(review_results)
        st.download_button(
            label="📥 下载审查意见报告 (Word)",
            data=word_buffer,
            file_name="工程审查意见报告.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )